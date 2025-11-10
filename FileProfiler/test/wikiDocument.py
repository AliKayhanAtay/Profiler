import os
import re
import requests
from docx import Document

WIKI_API_URL = "https://en.wikipedia.org/w/api.php"

HEADERS = {
    # Put something that identifies you (a nickname + email or site)
    "User-Agent": "KayhanWikiBot/0.1 (https://example.com; kayhan@example.com)"
}


def get_random_title():
    """Get a random article title from English Wikipedia."""
    params = {
        "action": "query",
        "format": "json",
        "list": "random",
        "rnnamespace": 0,   # main/article namespace
        "rnlimit": 1,
    }
    r = requests.get(WIKI_API_URL, params=params, headers=HEADERS)
    r.raise_for_status()
    data = r.json()
    return data["query"]["random"][0]["title"]


def get_page_text(title):
    """Get plain-text content of a Wikipedia page by title."""
    params = {
        "action": "query",
        "format": "json",
        "prop": "extracts",
        "explaintext": 1,
        "exsectionformat": "plain",
        "titles": title,
    }
    r = requests.get(WIKI_API_URL, params=params, headers=HEADERS)
    r.raise_for_status()
    data = r.json()
    pages = data["query"]["pages"]
    page = next(iter(pages.values()))
    return page.get("extract", "")


def safe_filename_from_title(title: str) -> str:
    """Convert a wiki title to a safe filename."""
    fname = title.replace(" ", "_")
    fname = re.sub(r"[^A-Za-z0-9_\-]", "_", fname)
    return fname or "article"


def save_random_articles_as_separate_docs(n_articles: int, out_dir: str = "wiki_docs"):
    out_dir = os.path.join(os.getcwd(), "FileProfiler", "test", "wiki_docs")
    os.makedirs(out_dir, exist_ok=True)

    for i in range(n_articles):
        title = get_random_title()
        text = get_page_text(title)
        print(f"[{i+1}/{n_articles}] Downloaded: {title}")

        doc = Document()
        doc.add_heading(title, level=1)

        for line in text.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

        fname = safe_filename_from_title(title) + ".docx"
        path = os.path.join(out_dir, fname)
        doc.save(path)
        print(f"  → Saved: {path}")


if __name__ == "__main__":
    save_random_articles_as_separate_docs(5, "wiki_random_docs")
